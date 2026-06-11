"""Insert Copernicus ensemble-sensitivity results into the DE/EN working papers."""

from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
ENSEMBLE_SUMMARY = REPO / "outputs" / "ensemble_rankings" / "copernicus2100_de_fr_top2_summary.csv"
COUNTRY_SENSITIVITY = REPO / "outputs" / "sensitivity" / "country_set_top2_summary.csv"
APPENDIX_TOP10 = REPO / "outputs" / "appendix" / "primary_top10.csv"
APPENDIX_WEIGHTED = REPO / "outputs" / "appendix" / "country_weighted_top2.csv"
APPENDIX_CRITERIA = REPO / "outputs" / "appendix" / "ranking_criteria_top2.csv"
CRITERIA_HEATMAP = REPO / "outputs" / "figures" / "ranking_criteria_top2_heatmap_de_fr.png"
TOP10_RANK_CURVE = REPO / "outputs" / "figures" / "top10_rank_curve_de_fr.png"
METHOD_FLOW = REPO / "outputs" / "figures" / "method_flow_diagram.png"
COUNTRY_MASK_HEATMAP = REPO / "outputs" / "figures" / "country_mask_top2_heatmap.png"
ENSEMBLE_DOTPLOT = REPO / "outputs" / "figures" / "ensemble_top2_dotplot.png"


def main(argv: list[str] | None = None) -> None:
    args = parse_args(argv)
    summary = pd.read_csv(args.ensemble_summary)
    country_sensitivity = pd.read_csv(args.country_sensitivity)
    appendix_top10 = pd.read_csv(args.appendix_top10)
    appendix_weighted = pd.read_csv(args.appendix_weighted)
    appendix_criteria = pd.read_csv(args.appendix_criteria)
    update_de(
        args.paper_dir / "heatwave_definition_working_paper_de_revised.docx",
        summary,
        country_sensitivity,
        appendix_top10,
        appendix_weighted,
        appendix_criteria,
        args.criteria_heatmap,
        args.top10_rank_curve,
        args.method_flow,
        args.country_mask_heatmap,
        args.ensemble_dotplot,
    )
    update_en(
        args.paper_dir / "heatwave_definition_working_paper_en_revised.docx",
        summary,
        country_sensitivity,
        appendix_top10,
        appendix_weighted,
        appendix_criteria,
        args.criteria_heatmap,
        args.top10_rank_curve,
        args.method_flow,
        args.country_mask_heatmap,
        args.ensemble_dotplot,
    )


def update_de(
    path: Path,
    summary: pd.DataFrame,
    country_sensitivity: pd.DataFrame,
    appendix_top10: pd.DataFrame,
    appendix_weighted: pd.DataFrame,
    appendix_criteria: pd.DataFrame,
    criteria_heatmap: Path,
    top10_rank_curve: Path,
    method_flow: Path,
    country_mask_heatmap: Path,
    ensemble_dotplot: Path,
) -> None:
    doc = Document(path)
    replace_text(
        doc,
        {
            "Version: überarbeitete Fassung, 1. Juni 2026":
            "Version: Zwischenstand, 2. Juni 2026",
            "Aus der vorliegenden Auswertung ergibt sich RCP4.5 / 2043 als primäres zukünftiges Hitzewellen-Szenario":
            "Aus der vorliegenden Auswertung ergibt sich RCP4.5 / IPSL-WRF / 2043 als primäres zukünftiges Hitzewellen-Szenario",
            "Für ein realistisches und zugleich belastendes zukünftiges Stresstest-Szenario wird RCP4.5 / 2043 als primäres Hitzewellen-Wetterjahr festgelegt. Die historisch":
            "Für ein realistisches und zugleich belastendes zukünftiges Stresstest-Szenario wird RCP4.5 / IPSL-WRF / 2043 als primäres Hitzewellen-Wetterjahr festgelegt. Die historisch",
            "RCP4.5 / 2070 eignet sich als Sensitivität":
            "RCP4.5 / IPSL-WRF / 2070 eignet sich als Sensitivität",
            "RCP8.5 / 2092 ist als sehr extremes Langfrist-Szenario":
            "RCP8.5 / MPI-CLM / 2092 ist als sehr extremes Langfrist-Szenario",
            "Historie, RCP4.5 und RCP8.5":
            "Historie, RCP4.5/IPSL-WRF und RCP8.5/MPI-CLM",
        },
    )
    insert_literature_gap_paragraph(
        doc,
        heading="2. Indexwahl",
        marker="Neuere Literatur bestätigt",
        text=(
            "Neuere Literatur bestätigt, dass Hitzewellendefinitionen weiterhin "
            "nicht vollständig standardisiert sind und dass die Wahl der Metrik "
            "räumliche Muster und Risikoeinschätzungen spürbar beeinflusst "
            "(Bunting et al., 2024; Schwingshackl et al., 2024). Zugleich zeigen "
            "aktuelle Energiesystemstudien, dass Wetterjahrauswahl und "
            "interannuelle Wettervariabilität zentrale Treiber robuster "
            "Systemauslegung sind (Gøtske et al., 2024; Antonini et al., 2024). "
            "Die hier adressierte Lücke liegt daher nicht in einer neuen "
            "Hitzewellenmetrik, sondern in einer transparenten, reproduzierbaren "
            "Definition konkreter Hitzewellen-Stresstestjahre für "
            "Energiesystemanalysen."
        ),
    )
    insert_method_flow_figure(
        doc,
        image_path=method_flow,
        caption=(
            "Abbildung 1: Methodischer Ablauf von den Temperaturdaten zur "
            "Hitzewellen-Szenarioauswahl."
        ),
        method_anchor_text="Ranke Szenariojahre nach der Summe der gitterzellbasierten HWMId-Jahreswerte in Deutschland und Frankreich.",
    )

    data_table = doc.tables[0]
    data_table.cell(2, 0).text = "Copernicus/CORDEX tasAdjust, RCP4.5 (IPSL-WRF; MPI/NCC als Sensitivität)"
    data_table.cell(3, 0).text = "Copernicus/CORDEX tasAdjust, RCP8.5 (MPI-CLM; IPSL/CNRM/NCC als Sensitivität)"

    result_table = doc.tables[1]
    result_table.cell(2, 0).text = "RCP4.5 / IPSL-WRF"
    result_table.cell(3, 0).text = "RCP8.5 / MPI-CLM"
    result_table.cell(3, 3).text = (
        "Extremere Langfrist-Sensitivität auf Basis der MPI-CLM-Modellkette."
    )

    insert_ensemble_section(
        doc,
        anchor_table=result_table,
        summary=summary,
        heading="Ensemble-Sensitivität der Copernicus-Rohdaten",
        note=(
            "Eine zusätzliche Rohdatenauswertung der lokal verfügbaren Copernicus2100-"
            "Ensembles zeigt, dass die stärksten Jahre deutlich von der gewählten "
            "GCM/RCM-Modellkette abhängen. Die Sensitivität dient zur Einordnung "
            "der Hauptauswahl; Szenariojahre werden daher mit Emissionspfad und "
            "GCM/RCM-Kombination berichtet."
        ),
        headers=[
            "Modellkette / Szenario",
            "Rang 1",
            "HWMId-Summe",
            "Rang 2",
            "HWMId-Summe",
        ],
        number_formatter=fmt_de,
    )
    insert_figure_after_table(
        doc,
        anchor_table=result_table,
        image_path=top10_rank_curve,
        caption=(
            "Abbildung 2: Top-10-Ranking der Hitzewellenjahre nach "
            "summierter gitterzellbasierter HWMId über Deutschland+Frankreich."
        ),
        remove_prefixes=(
            "Abbildung 2: Top-10-Ranking",
            "Abbildung 3: Top-10-Ranking",
            "Figure 2: Top-10 ranking",
            "Figure 3: Top-10 ranking",
        ),
        width=6.2,
    )
    renumber_spatial_and_criteria_captions(doc, german=True)
    insert_country_sensitivity_section(
        doc,
        anchor_table=doc.tables[-1],
        sensitivity=country_sensitivity,
        heading="Sensitivität der Ländermaske",
        note=(
            "Ergänzend zur Basismaske Deutschland+Frankreich wurde das Ranking "
            "für einzelne Länder und größere europäische Masken wiederholt. "
            "Die Top-2-Jahre bleiben für RCP4.5 in größeren Masken stabil, "
            "während die RCP8.5/MPI-CLM-Auswahl bei breiterer Abdeckung stärker "
            "in Richtung 2098/2092 kippt."
        ),
        headers=["Datenbasis", "Maske", "Rang 1", "Rang 2"],
        number_formatter=fmt_de,
    )
    insert_robustness_section(
        doc,
        anchor_table=doc.tables[3] if len(doc.tables) > 3 else doc.tables[-1],
        heading="Robustheit der Szenarioauswahl",
        text=(
            "Die Robustheitsprüfungen zeigen, dass die HWMId-basierte Hauptauswahl "
            "gegenüber einer Flächengewichtung stabil ist: HWMId-Summe, "
            "flächengewichteter HWMId-Mittelwert und ungewichteter HWMId-Mittelwert "
            "liefern für Deutschland+Frankreich dieselben Top-2-Jahre. Kriterien, "
            "die andere Ereigniseigenschaften betonen, verschieben einzelne "
            "Sensitivitätsjahre, insbesondere die maximale Einzelzellenintensität, "
            "die flächengewichtete Ereignisdauer und die jährliche Tmax-Anomalie. "
            "Die HWMId-Summe über Deutschland+Frankreich bleibt deshalb das "
            "Basiskriterium; alternative Kriterien werden als Sensitivität berichtet."
        ),
        image_path=criteria_heatmap,
        caption=(
            "Abbildung 4: Top-2-Jahre nach alternativen Ranking-Kriterien für "
            "Deutschland+Frankreich; Rang 2 steht in Klammern."
        ),
    )

    add_limitation(
        doc,
        "Die RCP8.5-Jahresauswahl ist modellkettenabhängig; daher müssen "
        "Emissionspfad und GCM/RCM-Kombination gemeinsam berichtet werden.",
    )
    insert_appendix_sections(
        doc,
        top10=appendix_top10,
        weighted=appendix_weighted,
        heading="Anhang: Sensitivitätstabellen",
        top10_note=(
            "Tabelle A1 zeigt das Top-10-Ranking für die drei im Haupttext "
            "verwendeten Datenbasen. Tabelle A2 zeigt die TYNDP-2024-PEMMDB-NT2040-gewichtete "
            "Sensitivität mit Kapazitäts- und Technologiegewichten."
        ),
        top10_headers=["Datenbasis", "Rang", "Jahr", "HWMId-Summe"],
        weighted_headers=["Datenbasis", "Gewichtung", "Rang 1", "Rang 2"],
        acknowledgement_heading="Danksagung",
        number_formatter=fmt_de,
        german=True,
    )
    insert_criteria_appendix(
        doc,
        criteria=appendix_criteria,
        acknowledgement_heading="Danksagung",
        note="Tabelle A3: Top-2-Auswahl nach alternativen Ranking-Kriterien.",
        headers=["Datenbasis", "Kriterium", "Rang 1", "Rang 2"],
        number_formatter=fmt_de,
        german=True,
    )
    insert_appendix_figures(
        doc,
        acknowledgement_heading="Danksagung",
        figures=[
            (
                country_mask_heatmap,
                "Abbildung A1: Top-2-Jahre nach Ländermaske und Datenbasis; Rang 2 steht in Klammern.",
                ("Abbildung A1:", "Figure A1:"),
            ),
            (
                ensemble_dotplot,
                "Abbildung A2: Ensemble-Sensitivität der Copernicus-Rohdaten; Punkte zeigen Rang 1, Dreiecke Rang 2.",
                ("Abbildung A2:", "Figure A2:"),
            ),
        ],
    )
    reorder_results_section(doc, german=True)
    insert_references(doc, "Literatur")
    mark_new_literature_red(doc)
    doc.save(path)


def update_en(
    path: Path,
    summary: pd.DataFrame,
    country_sensitivity: pd.DataFrame,
    appendix_top10: pd.DataFrame,
    appendix_weighted: pd.DataFrame,
    appendix_criteria: pd.DataFrame,
    criteria_heatmap: Path,
    top10_rank_curve: Path,
    method_flow: Path,
    country_mask_heatmap: Path,
    ensemble_dotplot: Path,
) -> None:
    doc = Document(path)
    replace_text(
        doc,
        {
            "Version: revised version, 1 June 2026":
            "Version: interim version, 2 June 2026",
            "The resulting primary future stress-test year is RCP4.5 / 2043":
            "The resulting primary future stress-test year is RCP4.5 / IPSL-WRF / 2043",
            "RCP4.5 / 2043 is selected as the primary future heatwave weather year":
            "RCP4.5 / IPSL-WRF / 2043 is selected as the primary future heatwave weather year",
            "RCP4.5 / 2070 can be used for a bandwidth sensitivity":
            "RCP4.5 / IPSL-WRF / 2070 can be used for a bandwidth sensitivity",
            "RCP8.5 / 2092 should be interpreted as a very extreme long-term case":
            "RCP8.5 / MPI-CLM / 2092 should be interpreted as a very extreme long-term case",
            "historical data, RCP4.5 and RCP8.5":
            "historical data, RCP4.5/IPSL-WRF and RCP8.5/MPI-CLM",
        },
    )
    insert_literature_gap_paragraph(
        doc,
        heading="2. Index choice",
        marker="Recent literature confirms",
        text=(
            "Recent literature confirms that heatwave definitions are still not "
            "fully standardized and that the choice of heat metric can materially "
            "affect spatial patterns and risk estimates (Bunting et al., 2024; "
            "Schwingshackl et al., 2024). At the same time, recent energy-system "
            "studies show that weather-year selection and interannual weather "
            "variability are central drivers of robust system design (Gøtske et al., "
            "2024; Antonini et al., 2024). The gap addressed here is therefore not "
            "a new heatwave metric, but a transparent and reproducible definition "
            "of concrete heatwave stress-test years for energy-system analysis."
        ),
    )
    insert_method_flow_figure(
        doc,
        image_path=method_flow,
        caption="Figure 1: Method workflow from temperature data to heatwave scenario selection.",
        method_anchor_text="Rank scenario years by summed grid-cell yearly HWMId in Germany and France.",
    )

    data_table = doc.tables[0]
    data_table.cell(2, 0).text = "Copernicus/CORDEX tasAdjust, RCP4.5 (IPSL-WRF; MPI/NCC sensitivity)"
    data_table.cell(3, 0).text = "Copernicus/CORDEX tasAdjust, RCP8.5 (MPI-CLM; IPSL/CNRM/NCC sensitivity)"

    result_table = doc.tables[1]
    result_table.cell(2, 0).text = "RCP4.5 / IPSL-WRF"
    result_table.cell(3, 0).text = "RCP8.5 / MPI-CLM"
    result_table.cell(3, 3).text = (
        "More extreme long-term sensitivity based on the MPI-CLM model chain."
    )

    insert_ensemble_section(
        doc,
        anchor_table=result_table,
        summary=summary,
        heading="Copernicus raw-data ensemble sensitivity",
        note=(
            "An additional raw-data run over the locally available Copernicus2100 "
            "ensembles shows that the strongest years depend strongly on the selected "
            "GCM/RCM model chain. The sensitivity run is used to contextualize the "
            "main selection; scenario years are therefore reported together with "
            "emission pathway and GCM/RCM combination."
        ),
        headers=[
            "Model chain / scenario",
            "Rank 1",
            "HWMId sum",
            "Rank 2",
            "HWMId sum",
        ],
        number_formatter=fmt_en,
    )
    insert_figure_after_table(
        doc,
        anchor_table=result_table,
        image_path=top10_rank_curve,
        caption=(
            "Figure 2: Top-10 ranking of heatwave years by summed grid-cell "
            "HWMId over Germany+France."
        ),
        remove_prefixes=(
            "Abbildung 2: Top-10-Ranking",
            "Abbildung 3: Top-10-Ranking",
            "Figure 2: Top-10 ranking",
            "Figure 3: Top-10 ranking",
        ),
        width=6.2,
    )
    renumber_spatial_and_criteria_captions(doc, german=False)
    insert_country_sensitivity_section(
        doc,
        anchor_table=doc.tables[-1],
        sensitivity=country_sensitivity,
        heading="Country-mask sensitivity",
        note=(
            "In addition to the Germany+France baseline mask, the ranking was "
            "repeated for individual countries and larger European masks. "
            "The RCP4.5 top years remain stable for broader masks, whereas the "
            "RCP8.5/MPI-CLM selection shifts towards 2098/2092 under wider coverage."
        ),
        headers=["Data basis", "Mask", "Rank 1", "Rank 2"],
        number_formatter=fmt_en,
    )
    insert_robustness_section(
        doc,
        anchor_table=doc.tables[3] if len(doc.tables) > 3 else doc.tables[-1],
        heading="Robustness of the scenario selection",
        text=(
            "The robustness checks show that the HWMId-based main selection is "
            "stable under area weighting: HWMId sum, area-weighted mean HWMId and "
            "unweighted mean HWMId return the same top-2 years for Germany+France. "
            "Criteria that emphasize different event properties shift individual "
            "sensitivity years, especially maximum grid-cell intensity, area-weighted "
            "event duration and annual Tmax anomaly. The HWMId sum over "
            "Germany+France is therefore retained as the baseline criterion, while "
            "alternative criteria are reported as sensitivities."
        ),
        image_path=criteria_heatmap,
        caption=(
            "Figure 4: Top-2 years by alternative ranking criterion for "
            "Germany+France; rank 2 is shown in parentheses."
        ),
    )

    add_limitation(
        doc,
        "The RCP8.5 year selection is model-chain dependent; emission pathway and "
        "GCM/RCM combination must therefore be reported together.",
    )
    insert_appendix_sections(
        doc,
        top10=appendix_top10,
        weighted=appendix_weighted,
        heading="Appendix: Sensitivity tables",
        top10_note=(
            "Table A1 lists the top-10 ranking for the three data bases used in "
            "the main text. Table A2 reports the TYNDP 2024 PEMMDB NT2040 weighted sensitivity "
            "using capacity and technology weights."
        ),
        top10_headers=["Data basis", "Rank", "Year", "HWMId sum"],
        weighted_headers=["Data basis", "Weighting", "Rank 1", "Rank 2"],
        acknowledgement_heading="Acknowledgements",
        number_formatter=fmt_en,
        german=False,
    )
    insert_criteria_appendix(
        doc,
        criteria=appendix_criteria,
        acknowledgement_heading="Acknowledgements",
        note="Table A3: Top-2 selection by alternative ranking criterion.",
        headers=["Data basis", "Criterion", "Rank 1", "Rank 2"],
        number_formatter=fmt_en,
        german=False,
    )
    insert_appendix_figures(
        doc,
        acknowledgement_heading="Acknowledgements",
        figures=[
            (
                country_mask_heatmap,
                "Figure A1: Top-2 years by country mask and data basis; rank 2 is shown in parentheses.",
                ("Abbildung A1:", "Figure A1:"),
            ),
            (
                ensemble_dotplot,
                "Figure A2: Copernicus raw-data ensemble sensitivity; circles show rank 1 and triangles rank 2.",
                ("Abbildung A2:", "Figure A2:"),
            ),
        ],
    )
    reorder_results_section(doc, german=False)
    insert_references(doc, "References")
    mark_new_literature_red(doc)
    doc.save(path)


def insert_ensemble_section(
    doc,
    anchor_table,
    summary: pd.DataFrame,
    heading: str,
    note: str,
    headers: list[str],
    number_formatter,
) -> None:
    if any(paragraph.text == heading for paragraph in doc.paragraphs):
        return

    heading_paragraph = doc.add_paragraph(heading, style="Heading 2")
    anchor_table._tbl.addnext(heading_paragraph._p)
    note_paragraph = doc.add_paragraph(note)
    heading_paragraph._p.addnext(note_paragraph._p)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    for ensemble, group in summary.groupby("ensemble", sort=False):
        top = group.sort_values("rank")
        if len(top) < 2:
            continue
        row = table.add_row().cells
        row[0].text = shorten_ensemble_label(str(ensemble))
        row[1].text = str(int(top.iloc[0]["year"]))
        row[2].text = number_formatter(float(top.iloc[0]["hwmid_sum"]))
        row[3].text = str(int(top.iloc[1]["year"]))
        row[4].text = number_formatter(float(top.iloc[1]["hwmid_sum"]))

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    note_paragraph._p.addnext(table._tbl)


def insert_country_sensitivity_section(
    doc,
    anchor_table,
    sensitivity: pd.DataFrame,
    heading: str,
    note: str,
    headers: list[str],
    number_formatter,
) -> None:
    if any(paragraph.text == heading for paragraph in doc.paragraphs):
        return

    heading_paragraph = doc.add_paragraph(heading, style="Heading 2")
    anchor_table._tbl.addnext(heading_paragraph._p)
    note_paragraph = doc.add_paragraph(note)
    heading_paragraph._p.addnext(note_paragraph._p)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    order = [
        "DE_FR",
        "DE_only",
        "FR_only",
        "DE_FR_Benelux_Alps",
        "Western_Central_Europe",
    ]
    dataset_order = ["Historical / E-OBS", "RCP4.5 / IPSL-WRF", "RCP8.5 / MPI-CLM"]
    sensitivity = sensitivity.copy()
    sensitivity["country_set"] = pd.Categorical(sensitivity["country_set"], order, ordered=True)
    sensitivity["dataset"] = pd.Categorical(sensitivity["dataset"], dataset_order, ordered=True)
    for (dataset, country_set), group in sensitivity.sort_values(["dataset", "country_set", "rank"]).groupby(
        ["dataset", "country_set"],
        observed=True,
    ):
        if len(group) < 2:
            continue
        top = group.sort_values("rank")
        row = table.add_row().cells
        row[0].text = localized_dataset_label(str(dataset), headers[0])
        row[1].text = mask_label(str(country_set))
        row[2].text = f"{int(top.iloc[0]['year'])} ({number_formatter(float(top.iloc[0]['hwmid_sum']))})"
        row[3].text = f"{int(top.iloc[1]['year'])} ({number_formatter(float(top.iloc[1]['hwmid_sum']))})"

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    note_paragraph._p.addnext(table._tbl)


def insert_robustness_section(
    doc,
    anchor_table,
    heading: str,
    text: str,
    image_path: Path,
    caption: str,
) -> None:
    if any(paragraph.text == heading for paragraph in doc.paragraphs):
        return

    heading_paragraph = doc.add_paragraph(heading, style="Heading 2")
    anchor_table._tbl.addnext(heading_paragraph._p)
    text_paragraph = doc.add_paragraph(text)
    heading_paragraph._p.addnext(text_paragraph._p)

    if image_path.exists():
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(str(image_path), width=Inches(6.1))
        text_paragraph._p.addnext(image_paragraph._p)

        caption_paragraph = doc.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph._p.addnext(caption_paragraph._p)


def insert_method_flow_figure(doc, image_path: Path, caption: str, method_anchor_text: str) -> None:
    remove_captioned_figure(doc, ("Abbildung 1: Methodischer Ablauf", "Figure 1: Method workflow"))
    if not image_path.exists():
        return
    anchor = find_paragraph_containing(doc, method_anchor_text)
    if anchor is None:
        anchor = find_paragraph(doc, "4. Methode") or find_paragraph(doc, "4. Method")
    if anchor is None:
        return
    image_paragraph, caption_paragraph = build_figure_paragraphs(doc, image_path, caption, width=6.2)
    anchor._p.addnext(image_paragraph._p)
    image_paragraph._p.addnext(caption_paragraph._p)


def insert_figure_after_table(
    doc,
    anchor_table,
    image_path: Path,
    caption: str,
    remove_prefixes: tuple[str, ...],
    width: float,
) -> None:
    remove_captioned_figure(doc, remove_prefixes)
    if not image_path.exists():
        return
    image_paragraph, caption_paragraph = build_figure_paragraphs(doc, image_path, caption, width=width)
    anchor_table._tbl.addnext(image_paragraph._p)
    image_paragraph._p.addnext(caption_paragraph._p)


def insert_appendix_figures(
    doc,
    acknowledgement_heading: str,
    figures: list[tuple[Path, str, tuple[str, ...]]],
) -> None:
    anchor = find_paragraph(doc, acknowledgement_heading)
    if anchor is None:
        return
    for _image_path, _caption, prefixes in figures:
        remove_captioned_figure(doc, prefixes)
    for image_path, caption, _prefixes in figures:
        if not image_path.exists():
            continue
        image_paragraph, caption_paragraph = build_figure_paragraphs(doc, image_path, caption, width=6.2)
        anchor._p.addprevious(image_paragraph._p)
        anchor._p.addprevious(caption_paragraph._p)


def build_figure_paragraphs(doc, image_path: Path, caption: str, width: float):
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_paragraph.runs:
        run.font.size = Pt(9)
    return image_paragraph, caption_paragraph


def remove_captioned_figure(doc, prefixes: tuple[str, ...]) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith(prefixes):
            parent = paragraph._p.getparent()
            previous = paragraph._p.getprevious()
            if previous is not None and ("<w:drawing" in previous.xml or "<w:pict" in previous.xml):
                parent.remove(previous)
            parent.remove(paragraph._p)


def renumber_spatial_and_criteria_captions(doc, german: bool) -> None:
    replacements = (
        (
            "Abbildung 1: Räumliche HWMId-Verteilung",
            "Abbildung 3: Räumliche HWMId-Verteilung",
        ),
        (
            "Abbildung 2: Räumliche HWMId-Verteilung",
            "Abbildung 3: Räumliche HWMId-Verteilung",
        ),
        (
            "Abbildung 2: Top-2-Jahre nach alternativen Ranking-Kriterien",
            "Abbildung 4: Top-2-Jahre nach alternativen Ranking-Kriterien",
        ),
        (
            "Figure 1: Spatial HWMId distribution",
            "Figure 3: Spatial HWMId distribution",
        ),
        (
            "Figure 2: Spatial HWMId distribution",
            "Figure 3: Spatial HWMId distribution",
        ),
        (
            "Figure 2: Top-2 years by alternative ranking criterion",
            "Figure 4: Top-2 years by alternative ranking criterion",
        ),
    )
    for paragraph in doc.paragraphs:
        for old, new in replacements:
            if paragraph.text.startswith(old):
                set_paragraph_text(paragraph, paragraph.text.replace(old, new, 1))


def reorder_results_section(doc, german: bool) -> None:
    if german:
        top10_caption = find_paragraph_starting(doc, ("Abbildung 2: Top-10-Ranking",))
        result_paragraph = find_paragraph_starting(
            doc,
            ("Für ein realistisches und zugleich belastendes zukünftiges Stresstest-Szenario",),
        )
        spatial_caption = find_paragraph_starting(doc, ("Abbildung 3: Räumliche HWMId-Verteilung",))
    else:
        top10_caption = find_paragraph_starting(doc, ("Figure 2: Top-10 ranking",))
        result_paragraph = find_paragraph_starting(
            doc,
            ("RCP4.5 / IPSL-WRF / 2043 is selected as the primary future heatwave weather year",),
        )
        spatial_caption = find_paragraph_starting(doc, ("Figure 3: Spatial HWMId distribution",))

    if top10_caption is None or result_paragraph is None or spatial_caption is None:
        return

    spatial_image = spatial_caption._p.getprevious()
    elements = [result_paragraph._p]
    if spatial_image is not None and ("<w:drawing" in spatial_image.xml or "<w:pict" in spatial_image.xml):
        elements.append(spatial_image)
    elements.append(spatial_caption._p)

    anchor = top10_caption._p
    for element in elements:
        anchor.addnext(element)
        anchor = element


def insert_appendix_sections(
    doc,
    top10: pd.DataFrame,
    weighted: pd.DataFrame,
    heading: str,
    top10_note: str,
    top10_headers: list[str],
    weighted_headers: list[str],
    acknowledgement_heading: str,
    number_formatter,
    german: bool,
) -> None:
    if any(paragraph.text == heading for paragraph in doc.paragraphs):
        return

    anchor = find_paragraph(doc, acknowledgement_heading)
    if anchor is None:
        anchor = doc.paragraphs[-1]

    heading_paragraph = doc.add_paragraph(heading, style="Heading 1")
    anchor._p.addprevious(heading_paragraph._p)
    note_paragraph = doc.add_paragraph(top10_note)
    anchor._p.addprevious(note_paragraph._p)

    top10_table = doc.add_table(rows=1, cols=len(top10_headers))
    top10_table.style = "Table Grid"
    fill_header(top10_table, top10_headers)
    for _, row_data in top10.iterrows():
        row = top10_table.add_row().cells
        row[0].text = localized_dataset(row_data["dataset"], german)
        row[1].text = str(int(row_data["rank"]))
        row[2].text = str(int(row_data["year"]))
        row[3].text = number_formatter(float(row_data["hwmid_sum"]))
    format_table(top10_table, left_columns={0})
    anchor._p.addprevious(top10_table._tbl)

    weighted_note = doc.add_paragraph(
        "Tabelle A2: TYNDP-2024-PEMMDB-NT2040-gewichtete Top-2-Auswahl."
        if german
        else "Table A2: TYNDP 2024 PEMMDB NT2040 weighted top-2 selection."
    )
    anchor._p.addprevious(weighted_note._p)

    weighted_table = doc.add_table(rows=1, cols=len(weighted_headers))
    weighted_table.style = "Table Grid"
    fill_header(weighted_table, weighted_headers)
    for (dataset, weighting), group in weighted.groupby(["dataset", "weighting"], sort=False):
        group = group.sort_values("rank")
        if len(group) < 2:
            continue
        row = weighted_table.add_row().cells
        row[0].text = localized_dataset(dataset, german)
        row[1].text = localized_weighting(weighting, german)
        row[2].text = f"{int(group.iloc[0]['year'])} ({number_formatter(float(group.iloc[0]['weighted_hwmid']))})"
        row[3].text = f"{int(group.iloc[1]['year'])} ({number_formatter(float(group.iloc[1]['weighted_hwmid']))})"
    format_table(weighted_table, left_columns={0, 1})
    anchor._p.addprevious(weighted_table._tbl)


def insert_criteria_appendix(
    doc,
    criteria: pd.DataFrame,
    acknowledgement_heading: str,
    note: str,
    headers: list[str],
    number_formatter,
    german: bool,
) -> None:
    existing = find_paragraph(doc, note)
    if existing is not None:
        following = existing._p.getnext()
        parent = existing._p.getparent()
        if following is not None and following.tag.endswith("tbl"):
            parent.remove(following)
        parent.remove(existing._p)

    anchor = find_paragraph(doc, acknowledgement_heading)
    if anchor is None:
        anchor = doc.paragraphs[-1]

    note_paragraph = doc.add_paragraph(note)
    anchor._p.addprevious(note_paragraph._p)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    fill_header(table, headers)
    for (dataset, criterion), group in criteria.groupby(["dataset", "criterion_label"], sort=False):
        group = group.sort_values("rank")
        if len(group) < 2:
            continue
        row = table.add_row().cells
        row[0].text = localized_dataset(dataset, german)
        row[1].text = localized_criterion(criterion, german)
        row[2].text = f"{int(group.iloc[0]['year'])} ({fmt_score(float(group.iloc[0]['score']), german)})"
        row[3].text = f"{int(group.iloc[1]['year'])} ({fmt_score(float(group.iloc[1]['score']), german)})"
    format_table(table, left_columns={0, 1})
    anchor._p.addprevious(table._tbl)


def insert_literature_gap_paragraph(doc, heading: str, marker: str, text: str) -> None:
    if any(marker in paragraph.text for paragraph in doc.paragraphs):
        return
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading and idx + 1 < len(doc.paragraphs):
            new_paragraph = doc.add_paragraph(text)
            doc.paragraphs[idx + 1]._p.addnext(new_paragraph._p)
            return


def insert_references(doc, heading: str) -> None:
    references = [
        "Antonini, E. G. A.; Di Bella, A.; Savelli, I.; Drouet, L.; Tavoni, M. (2024): Weather- and climate-driven power supply and demand time series for power and energy system analyses. Scientific Data 11, 1324. DOI: 10.1038/s41597-024-04129-8.",
        "Bunting, E. L.; Tolmanov, V.; Keellings, D. (2024): What is a heat wave: A survey and literature synthesis of heat wave definitions across the United States. PLOS Climate 3(9), e0000468. DOI: 10.1371/journal.pclm.0000468.",
        "Famooss Paolini, L.; Pascale, S.; Ruggieri, P. et al. (2026): Drivers of summer extreme temperature trends in Europe: insight from three major heat waves using flow analogues. Climate Dynamics 64, 235. DOI: 10.1007/s00382-026-08171-7.",
        "Gøtske, E. K.; Andresen, G. B.; Neumann, F.; Victoria, M. (2024): Designing a sector-coupled European energy system robust to 60 years of historical weather data. Nature Communications 15, 10680. DOI: 10.1038/s41467-024-54853-3.",
        "Schwingshackl, C.; Daloz, A. S.; Iles, C.; Aunan, K.; Sillmann, J. (2024): High-resolution projections of ambient heat for major European cities using different heat metrics. Natural Hazards and Earth System Sciences 24, 331-354. DOI: 10.5194/nhess-24-331-2024.",
        "Skinner, C. B.; Touma, D.; Barlow, M.; Singh, D.; King, T. (2025): The spatial extent of heat waves has changed over the past four decades. Communications Earth & Environment 6, 662. DOI: 10.1038/s43247-025-02661-y.",
    ]
    existing_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    anchor = find_paragraph(doc, heading)
    if anchor is None:
        return
    insertion_point = doc.paragraphs[-1]
    for paragraph in doc.paragraphs:
        if paragraph._p.getprevious() is anchor._p or paragraph.text.startswith(("Becker,", "Blazejczyk,", "Budd,", "Copernicus", "Cornes,", "IPCC", "Nairn,", "Perkins,", "Russo,", "WMO")):
            insertion_point = paragraph
    for reference in references:
        if reference.split(":", 1)[0] in existing_text:
            continue
        new_paragraph = doc.add_paragraph(reference)
        insertion_point._p.addnext(new_paragraph._p)
        insertion_point = new_paragraph


def mark_new_literature_red(doc) -> None:
    new_markers = (
        "Neuere Literatur bestätigt",
        "Recent literature confirms",
        "Antonini, E. G. A.",
        "Bunting, E. L.",
        "Famooss Paolini, L.",
        "Gøtske, E. K.",
        "Schwingshackl, C.",
        "Skinner, C. B.",
    )
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(new_markers):
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)


def find_paragraph(doc, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def find_paragraph_containing(doc, text: str):
    for paragraph in doc.paragraphs:
        if text in paragraph.text:
            return paragraph
    return None


def find_paragraph_starting(doc, prefixes: tuple[str, ...]):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefixes):
            return paragraph
    return None


def fill_header(table, headers: list[str]) -> None:
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)


def format_table(table, left_columns: set[int]) -> None:
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in left_columns else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def replace_text(doc, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph_text(paragraph, new_text)


def add_limitation(doc, text: str) -> None:
    if any(paragraph.text == text for paragraph in doc.paragraphs):
        return
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() in {"7. Einschränkungen", "7. Limitations"}:
            target = doc.paragraphs[idx + 1]
            new_paragraph = doc.add_paragraph(text, style="List Bullet")
            target._p.addprevious(new_paragraph._p)
            return


def set_paragraph_text(paragraph, text: str) -> None:
    style = paragraph.style
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(text)
    paragraph.style = style


def shorten_ensemble_label(label: str) -> str:
    return (
        label.replace("CNRM-CERFACS-CNRM-CM5 / CNRM-ALADIN63", "CNRM-CM5 / ALADIN63")
        .replace("IPSL-IPSL-CM5A-MR / IPSL-WRF381P", "IPSL-CM5A-MR / WRF381P")
        .replace("MPI-M-MPI-ESM-LR / CLMcom-CCLM4-8-17", "MPI-ESM-LR / CCLM4-8-17")
        .replace("NCC-NorESM1-M / DMI-HIRHAM5", "NorESM1-M / HIRHAM5")
        .replace("RCP45", "RCP4.5")
        .replace("RCP85", "RCP8.5")
    )


def mask_label(label: str) -> str:
    return {
        "DE_FR": "DE+FR",
        "DE_only": "DE",
        "FR_only": "FR",
        "DE_FR_Benelux_Alps": "DE+FR+Benelux+Alps",
        "Western_Central_Europe": "Western/Central Europe",
    }.get(label, label)


def localized_dataset_label(label: str, first_header: str) -> str:
    if first_header == "Datenbasis" and label == "Historical / E-OBS":
        return "Historisch / E-OBS"
    return label


def localized_dataset(label: str, german: bool) -> str:
    if german and label == "Historical / E-OBS":
        return "Historisch / E-OBS"
    return str(label)


def localized_weighting(label: str, german: bool) -> str:
    replacements = {
        "capacity_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 capacity" if not german else "TYNDP 2024 PEMMDB NT2040 Gesamtleistung",
        "renewables_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 RES capacity" if not german else "TYNDP 2024 PEMMDB NT2040 RES-Leistung",
        "solar_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 solar" if not german else "TYNDP 2024 PEMMDB NT2040 Solar",
        "pv_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 PV incl. rooftop" if not german else "TYNDP 2024 PEMMDB NT2040 PV inkl. rooftop",
        "wind_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 wind" if not german else "TYNDP 2024 PEMMDB NT2040 Wind",
        "wind_onshore_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 onshore wind" if not german else "TYNDP 2024 PEMMDB NT2040 Wind onshore",
        "wind_offshore_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 offshore wind" if not german else "TYNDP 2024 PEMMDB NT2040 Wind offshore",
        "hydro_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 hydro excl. pumped" if not german else "TYNDP 2024 PEMMDB NT2040 Hydro ohne Pumpspeicher",
        "pumped_hydro_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 pumped hydro" if not german else "TYNDP 2024 PEMMDB NT2040 Pumpspeicher",
        "bio_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 bio/waste" if not german else "TYNDP 2024 PEMMDB NT2040 Bio/Abfall",
        "nuclear_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 nuclear" if not german else "TYNDP 2024 PEMMDB NT2040 Kernenergie",
        "storage_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 battery storage" if not german else "TYNDP 2024 PEMMDB NT2040 Batteriespeicher",
        "storage_total_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 battery + pumped hydro" if not german else "TYNDP 2024 PEMMDB NT2040 Batterie + Pumpspeicher",
        "thermal_tyndp2024_pemmdb_nt2040": "TYNDP 2024 PEMMDB NT2040 thermal" if not german else "TYNDP 2024 PEMMDB NT2040 Thermal",
    }
    return replacements.get(str(label), str(label))


def localized_criterion(label: str, german: bool) -> str:
    if not german:
        return str(label)
    replacements = {
        "HWMId sum": "HWMId-Summe",
        "Area-weighted HWMId mean": "Flächengewichteter HWMId-Mittelwert",
        "Mean HWMId": "Mittlerer HWMId",
        "Maximum grid-cell HWMId": "Maximaler Gitterzellen-HWMId",
        "Area-weighted heatwave duration": "Flächengewichtete Hitzewellendauer",
        "Area-weighted annual Tmax anomaly": "Flächengewichtete jährliche Tmax-Anomalie",
    }
    return replacements.get(str(label), str(label))


def fmt_de(value: float) -> str:
    return f"{value:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_en(value: float) -> str:
    return f"{value:,.0f}"


def fmt_score(value: float, german: bool) -> str:
    if abs(value) >= 1000:
        return fmt_de(value) if german else fmt_en(value)
    text = f"{value:.1f}"
    return text.replace(".", ",") if german else text


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--paper-dir", type=Path, required=True)
    parser.add_argument("--ensemble-summary", type=Path, default=ENSEMBLE_SUMMARY)
    parser.add_argument("--country-sensitivity", type=Path, default=COUNTRY_SENSITIVITY)
    parser.add_argument("--appendix-top10", type=Path, default=APPENDIX_TOP10)
    parser.add_argument("--appendix-weighted", type=Path, default=APPENDIX_WEIGHTED)
    parser.add_argument("--appendix-criteria", type=Path, default=APPENDIX_CRITERIA)
    parser.add_argument("--criteria-heatmap", type=Path, default=CRITERIA_HEATMAP)
    parser.add_argument("--top10-rank-curve", type=Path, default=TOP10_RANK_CURVE)
    parser.add_argument("--method-flow", type=Path, default=METHOD_FLOW)
    parser.add_argument("--country-mask-heatmap", type=Path, default=COUNTRY_MASK_HEATMAP)
    parser.add_argument("--ensemble-dotplot", type=Path, default=ENSEMBLE_DOTPLOT)
    return parser.parse_args(argv)


if __name__ == "__main__":
    main()
